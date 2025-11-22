from docx import Document


def add_title(document, title):
    document.add_heading(title, level=0)


def add_user(document, username):
    document.add_heading(username, level=1)


def add_question(document, question):
    p = document.add_paragraph()
    run = p.add_run(question)
    run.bold = True


def add_answer(document, answer):
    document.add_paragraph(answer)


if __name__ == "__main__":
    document = Document()

    add_title(document, 'Week 10: Refinement & iteration')

    add_user(document, 'Max Buster')

    add_question(document, 'What is the meaning of life?')
    add_answer(document, 'Answer: 42')

    document.save("output/example.docx")
