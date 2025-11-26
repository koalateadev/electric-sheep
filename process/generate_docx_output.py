import json
import os

from docx import Document

from process.map_inputs import generate_output
from process.parse_cohorts import parse_cohorts
from process.parse_responses import parse_responses


def _add_title(document, title):
    document.add_heading(title, level=0)


def _add_user(document, username):
    document.add_heading(username, level=1)


def _add_question(document, question):
    p = document.add_paragraph()
    run = p.add_run(question)
    run.bold = True


def _add_answer(document, answer):
    document.add_paragraph(answer)


def write_to_docx(results, base_path):
    for facilitator in results:
        for week in results[facilitator]:
            path = base_path + facilitator
            document = Document()

            _add_title(document, week)

            for user in results[facilitator][week]:
                _add_user(document, user)

                for response in results[facilitator][week][user]["responses"]:
                    _add_question(document, response["question"])
                    _add_answer(document, response["answer"])

            os.makedirs(path, mode=0o777, exist_ok=True)
            document.save(path + "/" + week + ".docx")


if __name__ == "__main__":
    base_path = os.path.dirname(os.path.abspath(__file__))

    results_file_name = base_path + "/input/responses.csv"
    cohorts_file_name = base_path + "/input/cohorts.json"
    student_id_file_name = base_path + "/input/airtable_name_to_teachable_name.json"
    output_file_name = base_path + "/output/output.json"

    parsed_input = parse_responses(results_file_name)
    parsed_mapping = parse_cohorts(cohorts_file_name)
    parsed_student_ids = parse_cohorts(student_id_file_name)

    results = generate_output(parsed_input, parsed_mapping, parsed_student_ids)

    write_to_docx(results, "./process/output/")
